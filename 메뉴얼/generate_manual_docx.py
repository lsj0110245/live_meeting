import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_manual_docx(input_md_path, output_docx_path):
    """
    Parses a Markdown file and creates a styled DOCX document.
    """
    if not os.path.exists(input_md_path):
        print(f"Error: Input file '{input_md_path}' not found.")
        return

    doc = Document()
    
    # Set default font to ensure Korean characters render correctly
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    
    # Configure headings style
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = 'Malgun Gothic'
        h_font.bold = True
        if i == 1:
            h_font.size = Pt(24)
            h_font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        elif i == 2:
            h_font.size = Pt(18)
            h_font.color.rgb = RGBColor(0, 102, 204) # Medium Blue
        elif i == 3:
            h_font.size = Pt(14)
            h_font.color.rgb = RGBColor(51, 51, 51) # Dark Gray

    with open(input_md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_content = []
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        line = line.strip()
        
        # 1. Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph('\n'.join(code_block_content))
                p.style = 'Quote' # Use Quote style for code blocks simpler
                runner = p.runs[0] if p.runs else p.add_run()
                runner.font.name = 'Consolas'
                runner.font.size = Pt(10)
                
                in_code_block = False
                code_block_content = []
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue

        # 2. Handle Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                # Prepare headers
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                table_headers = cells
            elif '---' in line:
                # Skip separator line
                continue
            else:
                # Add row
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                # If we have headers but haven't created the table yet, create it now
                if table_headers and not table_rows:
                     # Wait until we have rows or execute creation at end of block
                     pass
                table_rows.append(cells)
            continue
        else:
            if in_table:
                # End of table block, render table
                if table_headers:
                    table = doc.add_table(rows=1, cols=len(table_headers))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(table_headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    
                    for row_data in table_rows:
                        # Ensure row length matches header length
                        if len(row_data) != len(table_headers):
                            # Pad or truncate
                            row_data = row_data[:len(table_headers)] + [''] * (len(table_headers) - len(row_data))
                            
                        row_cells = table.add_row().cells
                        for i, cell_data in enumerate(row_data):
                            row_cells[i].text = cell_data
                
                in_table = False
                table_headers = []
                table_rows = []

        if not line:
            continue

        # 3. Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            
        # 4. Handle Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Ordered list
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            
        # 5. Handle Horizontal Rules
        elif line.startswith('---'):
            doc.add_paragraph('--------------------------------------------------')
            
        # 6. Normal Text
        else:
            p = doc.add_paragraph(line)
        
        # Apply strict bold formatting for **text**
        # (Simple implementation: iterate paragraphs and apply bold run if ** detected)
        # Note: Ideally this should be done during paragraph creation, but here we do a post-pass or just handle simple cases.
        # For simplicity in this script, we'll leave as text or do a simple replace if critical.
        
    doc.save(output_docx_path)
    print(f"Successfully created: {output_docx_path}")

if __name__ == "__main__":
    md_path = "c:/big20/live_meeting/MEETING_MODEL_MANUAL.md"
    docx_path = "c:/big20/live_meeting/MEETING_MODEL_MANUAL.docx"
    
    try:
        create_manual_docx(md_path, docx_path)
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        # Try to install python-docx if module not found
        import subprocess
        import sys
        if "No module named 'docx'" in str(e):
             print("Installing python-docx...")
             subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
             print("Retrying...")
             create_manual_docx(md_path, docx_path)
